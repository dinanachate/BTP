from course_build_agents.knowledge_retriever import KnowledgeRetrieverAgent
from course_build_agents.knowledge_enhancer import KnowledgeEnhancerAgent
from course_build_agents.course_generator import CourseGeneratorAgent
import json
import os
import sys
import logging
from datetime import datetime
from io import StringIO


class StreamingPrintCapture:
    """Captures print statements in a buffer for later retrieval."""
    def __init__(self):
        self.original_stdout = sys.stdout
        self.buffer = []

    def write(self, text):
        # Write to original stdout
        self.original_stdout.write(text)
        # Capture in buffer
        self.buffer.append(text)

    def flush(self):
        self.original_stdout.flush()

    def get_and_clear(self):
        """Get captured content and clear buffer."""
        content = ''.join(self.buffer)
        self.buffer = []
        return content


def stream_course_generation_progress(subject, config=None):
    """
    Stream course generation with progress updates.
    Yields progress messages that can be sent as reasoning_content.

    Yields:
        dict: {'type': 'progress'|'complete', 'content': str, 'results': dict}
    """
    config = config or {}

    # Setup print capture
    capture = StreamingPrintCapture()
    old_stdout = sys.stdout
    sys.stdout = capture

    try:
        retriever = KnowledgeRetrieverAgent(
            top_k_per_query=config.get('retriever_top_k', 5)
        )
        enhancer = KnowledgeEnhancerAgent(
            max_iterations=config.get('enhancer_iterations', 3),
            top_k=config.get('enhancer_top_k', 5)
        )
        course_generator = CourseGeneratorAgent()

        # Phase 1: Knowledge Retrieval
        print("=" * 80)
        print(f"GÉNÉRATION DE COURS MULTI-AGENTS")
        print(f"Sujet: {subject}")
        print(f"Démarré à: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("=" * 80)
        print("\n" + "=" * 80)
        print("PHASE 1: RÉCUPÉRATION DES CONNAISSANCES")
        print("=" * 80)

        # Yield header
        header = capture.get_and_clear()
        if header:
            yield {'type': 'progress', 'content': header}

        knowledge_base, sources = retriever.retrieve_knowledge(subject)

        # Yield retrieval logs
        retrieval_logs = capture.get_and_clear()
        if retrieval_logs:
            yield {'type': 'progress', 'content': retrieval_logs}

        # Phase 2: Knowledge Enhancement
        print("\n" + "=" * 80)
        print("PHASE 2: AMÉLIORATION DES CONNAISSANCES")
        print("=" * 80)

        phase2_header = capture.get_and_clear()
        if phase2_header:
            yield {'type': 'progress', 'content': phase2_header}

        enhanced_knowledge, all_sources = enhancer.enhance_knowledge(
            subject, knowledge_base, sources
        )

        # Yield enhancement logs
        enhancement_logs = capture.get_and_clear()
        if enhancement_logs:
            yield {'type': 'progress', 'content': enhancement_logs}

        sources_added = len(all_sources) - len(sources)

        # Phase 3: Course Generation
        print("\n" + "=" * 80)
        print("PHASE 3: GÉNÉRATION DE LA STRUCTURE DU COURS")
        print("=" * 80)

        phase3_header = capture.get_and_clear()
        if phase3_header:
            yield {'type': 'progress', 'content': phase3_header}

        course_structure = course_generator.generate_course(
            subject, enhanced_knowledge, all_sources
        )

        # Yield course generation logs
        generation_logs = capture.get_and_clear()
        if generation_logs:
            yield {'type': 'progress', 'content': generation_logs}

        # Generate markdown
        print("\n" + "=" * 80)
        print("GÉNÉRATION DU MARKDOWN")
        print("=" * 80)
        print(f"   Génération du contenu markdown...")

        markdown_header = capture.get_and_clear()
        if markdown_header:
            yield {'type': 'progress', 'content': markdown_header}

        course_markdown = course_generator.get_markdown_content()

        print(f"   ✓ Markdown généré ({len(course_markdown)} caractères)")
        print("\n" + "=" * 80)
        print("PROCESSUS TERMINÉ AVEC SUCCÈS")
        print(f"Chapitres: {course_structure.get('total_chapters', 0)}")
        print(f"Sources: {len(all_sources)} (dont {sources_added} ajoutées)")
        print("=" * 80)

        final_logs = capture.get_and_clear()
        if final_logs:
            yield {'type': 'progress', 'content': final_logs}

        # Return final results
        results = {
            'course_structure': course_structure,
            'course_markdown': course_markdown,
            'initial_source_count': len(sources),
            'final_source_count': len(all_sources),
            'sources_added': sources_added
        }

        yield {'type': 'complete', 'content': '', 'results': results}

    finally:
        # Restore stdout
        sys.stdout = old_stdout


class LogCapture:
    """Captures both print statements and logging to a file and memory."""

    def __init__(self, log_file_path):
        self.log_file_path = log_file_path
        self.log_buffer = StringIO()
        self.terminal = sys.stdout

        # Setup file logger
        self.file_handler = logging.FileHandler(log_file_path, mode='w', encoding='utf-8')
        self.file_handler.setLevel(logging.INFO)
        formatter = logging.Formatter('%(asctime)s - %(message)s', datefmt='%Y-%m-%d %H:%M:%S')
        self.file_handler.setFormatter(formatter)

        # Create logger
        self.logger = logging.getLogger('CourseGenerator')
        self.logger.setLevel(logging.INFO)
        self.logger.addHandler(self.file_handler)

    def write(self, message):
        """Intercept stdout writes."""
        self.terminal.write(message)
        self.log_buffer.write(message)
        # Also write to file, removing ANSI codes
        clean_message = self._remove_ansi(message)
        if clean_message.strip():
            self.logger.info(clean_message.rstrip())

    def flush(self):
        """Flush both outputs."""
        self.terminal.flush()
        self.log_buffer.flush()

    def _remove_ansi(self, text):
        """Remove ANSI escape codes from text."""
        import re
        ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
        return ansi_escape.sub('', text)

    def get_logs(self):
        """Get all captured logs as string."""
        return self.log_buffer.getvalue()

    def close(self):
        """Close file handler."""
        self.file_handler.close()
        self.logger.removeHandler(self.file_handler)


class MultiAgentOrchestratorWithLogging:
    """
    Enhanced orchestrator with logging and Word export capabilities.

    Flow:
    1. Knowledge Retriever: Gathers comprehensive knowledge
    2. Knowledge Enhancer: Identifies gaps and improves quality
    3. Course Generator: Creates structured course from knowledge
    4. Exports to Word document with proper formatting
    """

    def __init__(self, config=None):
        """
        Initialize orchestrator with optional configuration.

        config example:
        {
            'retriever_top_k': 5,
            'enhancer_iterations': 3,
            'enhancer_top_k': 5,
            'output_dir': './output',
            'enable_logging': True
        }
        """
        config = config or {}

        self.retriever = KnowledgeRetrieverAgent(
            top_k_per_query=config.get('retriever_top_k', 5)
        )
        self.enhancer = KnowledgeEnhancerAgent(
            max_iterations=config.get('enhancer_iterations', 3),
            top_k=config.get('enhancer_top_k', 5)
        )
        self.course_generator = CourseGeneratorAgent()

        self.output_dir = config.get('output_dir', './output')
        os.makedirs(self.output_dir, exist_ok=True)

        self.results = {}
        self.enable_logging = config.get('enable_logging', True)
        self.log_capture = None

    def run(self, subject):
        """
        Execute the complete multi-agent workflow with logging.

        Args:
            subject: The topic to create a course about

        Returns:
            dict: Complete results including knowledge base, course structure, and log file path
        """
        # Setup logging
        if self.enable_logging:
            log_file_path = os.path.join(self.output_dir, f'course_generation_log_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt')
            self.log_capture = LogCapture(log_file_path)
            old_stdout = sys.stdout
            sys.stdout = self.log_capture

        try:
            print("=" * 80)
            print(f"MULTI-AGENT COURSE GENERATION SYSTEM")
            print(f"Subject: {subject}")
            print(f"Started at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            print("=" * 80)

            start_time = datetime.now()

            # AGENT 1: Knowledge Retrieval
            print("\n" + "=" * 80)
            print("PHASE 1: KNOWLEDGE RETRIEVAL")
            print("=" * 80)
            knowledge_base, sources = self.retriever.retrieve_knowledge(subject)

            self.results['initial_knowledge'] = knowledge_base
            self.results['initial_sources'] = sources
            self.results['initial_source_count'] = len(sources)

            # Save initial knowledge
            self._save_knowledge(knowledge_base, sources, 'initial_knowledge.md')

            # AGENT 2: Knowledge Enhancement
            print("\n" + "=" * 80)
            print("PHASE 2: KNOWLEDGE ENHANCEMENT")
            print("=" * 80)
            enhanced_knowledge, all_sources = self.enhancer.enhance_knowledge(
                subject, knowledge_base, sources
            )

            self.results['enhanced_knowledge'] = enhanced_knowledge
            self.results['all_sources'] = all_sources
            self.results['final_source_count'] = len(all_sources)
            self.results['sources_added'] = len(all_sources) - len(sources)

            # Save enhanced knowledge
            self._save_knowledge(enhanced_knowledge, all_sources, 'enhanced_knowledge.md')

            # AGENT 3: Course Generation
            print("\n" + "=" * 80)
            print("PHASE 3: COURSE STRUCTURE GENERATION")
            print("=" * 80)
            course_structure = self.course_generator.generate_course(
                subject, enhanced_knowledge, all_sources
            )

            self.results['course_structure'] = course_structure

            # Generate markdown content
            print("\n" + "=" * 80)
            print("GENERATING COURSE MARKDOWN")
            print("=" * 80)
            course_markdown = self.course_generator.get_markdown_content()
            print(f"   Markdown content generated ({len(course_markdown)} characters)")

            # Optionally save markdown to file if logging is enabled
            if self.enable_logging:
                course_md_path = os.path.join(self.output_dir, 'course_structure.md')
                self.course_generator.export_to_markdown(course_md_path)

            # Summary
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()

            print("\n" + "=" * 80)
            print("PROCESS COMPLETED SUCCESSFULLY")
            print("=" * 80)
            print(f"Subject: {subject}")
            print(f"Initial sources: {self.results['initial_source_count']}")
            print(f"Sources added by enhancer: {self.results['sources_added']}")
            print(f"Total sources: {self.results['final_source_count']}")
            print(f"Total chapters: {course_structure.get('total_chapters', 0)}")
            print(f"Duration: {duration:.2f} seconds")
            if self.enable_logging:
                print(f"\nOutput directory: {self.output_dir}")
            print("=" * 80)

            # Add markdown content to results
            self.results['course_markdown'] = course_markdown

            return self.results

        finally:
            # Restore stdout and close logging
            if self.enable_logging:
                sys.stdout = old_stdout
                self.log_capture.close()

    def _save_knowledge(self, knowledge, sources, filename):
        """Save knowledge base with sources to markdown file."""
        filepath = os.path.join(self.output_dir, filename)

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"# Knowledge Base\n\n")
            f.write(knowledge)
            f.write("\n\n---\n\n")
            f.write("## Sources\n\n")
            for i, source in enumerate(sources, 1):
                f.write(f"{i}. [{source['title']}]({source['url']})\n")

        print(f"   Saved: {filepath}")

    def _save_json_results(self):
        """Save complete results as JSON for programmatic access."""
        filepath = os.path.join(self.output_dir, 'results.json')

        # Create simplified version (without full text for readability)
        json_results = {
            'initial_source_count': self.results['initial_source_count'],
            'final_source_count': self.results['final_source_count'],
            'sources_added': self.results['sources_added'],
            'course_structure': self.results['course_structure'],
            'sources': [
                {
                    'id': s['id'],
                    'title': s['title'],
                    'url': s['url']
                }
                for s in self.results['all_sources']
            ]
        }

        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(json_results, f, ensure_ascii=False, indent=2)

        print(f"   Saved: {filepath}")

    def _export_to_word(self, course_structure, sources, output_path):
        """Export course structure to a professionally formatted Word document."""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        print(f"   Generating Word document...")

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Title
        title = doc.add_heading(course_structure['course_title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Course metadata
        doc.add_paragraph()
        meta_para = doc.add_paragraph()
        meta_para.add_run('Description: ').bold = True
        meta_para.add_run(course_structure.get('description', ''))

        meta_para = doc.add_paragraph()
        meta_para.add_run('Public cible: ').bold = True
        meta_para.add_run(course_structure.get('target_audience', ''))

        meta_para = doc.add_paragraph()
        meta_para.add_run('Nombre de chapitres: ').bold = True
        meta_para.add_run(str(course_structure.get('total_chapters', 0)))

        doc.add_page_break()

        # Table of Contents
        doc.add_heading('Table des matières', level=1)
        for idx, chapter in enumerate(course_structure['chapters'], 1):
            toc_para = doc.add_paragraph(style='List Number')
            toc_para.add_run(chapter['title'])
            if 'subchapters' in chapter:
                for sub_idx, subchapter in enumerate(chapter['subchapters'], 1):
                    sub_para = doc.add_paragraph(style='List Bullet 2')
                    sub_para.add_run(f"{idx}.{sub_idx} {subchapter['title']}")

        doc.add_page_break()

        # Chapters
        for ch_idx, chapter in enumerate(course_structure['chapters'], 1):
            # Chapter heading
            doc.add_heading(f"Chapitre {ch_idx}: {chapter['title']}", level=1)

            # Chapter description
            doc.add_paragraph(chapter.get('description', ''))

            # Learning objectives
            if 'learning_objectives' in chapter and chapter['learning_objectives']:
                doc.add_heading('Objectifs d\'apprentissage', level=2)
                for obj in chapter['learning_objectives']:
                    doc.add_paragraph(obj, style='List Bullet')

            # Subchapters
            if 'subchapters' in chapter:
                for sub_idx, subchapter in enumerate(chapter['subchapters'], 1):
                    # Subchapter heading
                    doc.add_heading(f"{ch_idx}.{sub_idx} - {subchapter['title']}", level=2)

                    # Content to cover
                    if 'content_to_cover' in subchapter and subchapter['content_to_cover']:
                        doc.add_heading('Contenu à couvrir:', level=3)
                        for content in subchapter['content_to_cover']:
                            doc.add_paragraph(content, style='List Bullet')

                    # Practical elements
                    if 'practical_elements' in subchapter and subchapter['practical_elements']:
                        doc.add_heading('Éléments pratiques:', level=3)
                        for element in subchapter['practical_elements']:
                            doc.add_paragraph(element, style='List Bullet')

            doc.add_page_break()

        # Sources/References
        doc.add_heading('Sources et références', level=1)
        for idx, source in enumerate(sources, 1):
            source_para = doc.add_paragraph(style='List Number')
            source_para.add_run(f"{source['title']}")
            if source.get('url'):
                source_para.add_run(f"\n   URL: {source['url']}")

        # Save document
        doc.save(output_path)
        print(f"   Course Word document saved: {output_path}")


def main():
    """Example usage of the enhanced multi-agent system."""

    # Configuration
    config = {
        'retriever_top_k': 5,          # Number of sources per query in retrieval
        'enhancer_iterations': 3,       # Max iterations for knowledge enhancement
        'enhancer_top_k': 5,            # Sources per gap-filling query
        'output_dir': './course_output',
        'enable_logging': True
    }

    # Initialize orchestrator
    orchestrator = MultiAgentOrchestratorWithLogging(config)

    # Run the system
    subject = "Machine Learning"  # Change this to your subject
    results = orchestrator.run(subject)

    print("\n✅ Course generation complete!")
    print(f"📁 Check the output in: {config['output_dir']}/")
    print(f"📄 Word document: {results.get('course_docx_path')}")
    print(f"📋 Log file: {results.get('log_file_path')}")


if __name__ == "__main__":
    main()
